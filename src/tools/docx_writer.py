"""
Ferramenta para criar playbooks DOCX com identidade visual Lector
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from pathlib import Path
from typing import List, Dict, Any, Optional
from PIL import Image as PILImage
import io

from ..config import BrandIdentity


class DocxWriter:
    """Cria playbooks DOCX formatados com identidade visual Lector"""

    def __init__(self, brand: BrandIdentity, output_path: Path):
        self.brand = brand
        self.output_path = Path(output_path)
        self.doc = Document()
        self._setup_styles()

    def _hex_to_rgb(self, hex_color: str) -> RGBColor:
        """Converte hex para RGBColor"""
        hex_color = hex_color.lstrip("#")
        return RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16)
        )

    def _setup_styles(self):
        """Configura estilos do documento com identidade Lector"""

        # Configurar margens
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Estilo Título Principal
        style = self.doc.styles["Title"]
        font = style.font
        font.name = self.brand.heading_font
        font.size = Pt(self.brand.title_size)
        font.bold = True
        font.color.rgb = self._hex_to_rgb(self.brand.primary_color)

        # Estilos de Título
        for i in range(1, 4):
            style_name = f"Heading {i}"
            style = self.doc.styles[style_name]
            font = style.font
            font.name = self.brand.heading_font
            font.bold = True

            if i == 1:
                font.size = Pt(self.brand.heading1_size)
                font.color.rgb = self._hex_to_rgb(self.brand.primary_color)
            elif i == 2:
                font.size = Pt(self.brand.heading2_size)
                font.color.rgb = self._hex_to_rgb(self.brand.secondary_color)
            else:
                font.size = Pt(self.brand.heading3_size)
                font.color.rgb = self._hex_to_rgb(self.brand.text_color)

        # Estilo Normal
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self.brand.body_font
        font.size = Pt(self.brand.body_size)
        font.color.rgb = self._hex_to_rgb(self.brand.text_color)

        # Configurar parágrafo normal
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = self.brand.line_spacing
        paragraph_format.space_after = Pt(self.brand.paragraph_spacing)

    def add_title(self, text: str):
        """Adiciona título principal"""
        title = self.doc.add_heading(text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adicionar linha decorativa após título
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("─" * 50)
        run.font.color.rgb = self._hex_to_rgb(self.brand.secondary_color)
        run.font.size = Pt(8)

    def add_heading(self, text: str, level: int = 1):
        """Adiciona título de seção"""
        if level < 1:
            level = 1
        if level > 3:
            level = 3

        heading = self.doc.add_heading(text, level=level)

        # Adicionar espaço extra antes de headings principais
        if level == 1:
            heading.paragraph_format.space_before = Pt(self.brand.section_spacing)

    def add_paragraph(self, text: str, style: Optional[str] = None):
        """Adiciona parágrafo formatado"""
        if style:
            para = self.doc.add_paragraph(text, style=style)
        else:
            para = self.doc.add_paragraph(text)

        return para

    def add_bullet_list(self, items: List[str]):
        """Adiciona lista com marcadores"""
        for item in items:
            para = self.doc.add_paragraph(item, style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.25)

    def add_numbered_list(self, items: List[str]):
        """Adiciona lista numerada"""
        for item in items:
            para = self.doc.add_paragraph(item, style="List Number")
            para.paragraph_format.left_indent = Inches(0.25)

    def add_info_box(self, title: str, content: str, box_type: str = "info"):
        """Adiciona caixa de informação estilizada"""
        # Determinar cor baseada no tipo
        if box_type == "warning":
            color = self.brand.accent_color
            icon = "⚠️"
        elif box_type == "tip":
            color = self.brand.secondary_color
            icon = "💡"
        else:
            color = self.brand.primary_color
            icon = "ℹ️"

        # Criar parágrafo com borda simulada
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.2)
        para.paragraph_format.right_indent = Inches(0.2)

        # Ícone e título
        run = para.add_run(f"{icon} {title}")
        run.bold = True
        run.font.color.rgb = self._hex_to_rgb(color)

        # Conteúdo
        content_para = self.doc.add_paragraph(content)
        content_para.paragraph_format.left_indent = Inches(0.4)
        content_para.paragraph_format.space_after = Pt(12)

    def add_step(self, number: int, title: str, description: str):
        """Adiciona um passo do playbook formatado"""
        # Número do passo em destaque
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)

        # Círculo com número
        run = para.add_run(f"PASSO {number}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = self._hex_to_rgb(self.brand.primary_color)

        # Título do passo
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = self._hex_to_rgb(self.brand.text_color)
        title_para.paragraph_format.space_after = Pt(6)

        # Descrição
        if description:
            desc_para = self.doc.add_paragraph(description)
            desc_para.paragraph_format.left_indent = Inches(0.2)
            desc_para.paragraph_format.space_after = Pt(12)

    def add_table(self, data: List[List[str]]):
        """Adiciona tabela estilizada"""
        if not data or len(data) < 1:
            return

        table = self.doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = "Table Grid"

        # Preencher dados
        for i, row in enumerate(data):
            for j, cell_text in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = cell_text

                # Estilizar header
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = self._hex_to_rgb("#FFFFFF")
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Cor de fundo
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    cell._tc.get_or_add_tcPr().append(parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{self.brand.primary_color.lstrip("#")}"></w:shd>'
                    ))

    def add_image(self, image_data: bytes, width: Optional[float] = None):
        """Adiciona imagem ao documento"""
        try:
            image_stream = io.BytesIO(image_data)

            if width:
                self.doc.add_picture(image_stream, width=Inches(width))
            else:
                self.doc.add_picture(image_stream, width=Inches(5.5))

        except Exception as e:
            print(f"Erro ao adicionar imagem: {e}")

    def add_separator(self):
        """Adiciona separador visual"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run("• • •")
        run.font.color.rgb = self._hex_to_rgb(self.brand.secondary_color)

    def add_page_break(self):
        """Adiciona quebra de página"""
        self.doc.add_page_break()

    def save(self):
        """Salva o documento"""
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(self.output_path)
        return self.output_path
