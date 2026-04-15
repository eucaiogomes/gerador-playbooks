#!/usr/bin/env python3
"""
Lector Playbook Converter - Versão Simplificada
Converte documentos DOCX em playbooks formatados com identidade visual Lector
"""

import os
import sys
import re
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass

# Configurar caminho para imports
sys.path.insert(0, str(Path(__file__).parent / "src"))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from PIL import Image as PILImage
except ImportError as e:
    print(f"Erro ao importar bibliotecas: {e}")
    print("Instalando dependencias...")
    os.system("pip install python-docx pillow")
    sys.exit(1)


@dataclass
class BrandIdentity:
    """Identidade visual Lector"""
    primary_color: str = "#1E3A5F"
    secondary_color: str = "#2563EB"
    accent_color: str = "#D97757"
    background_color: str = "#FFFFFF"
    text_color: str = "#2D3748"
    light_gray: str = "#F7FAFC"
    border_color: str = "#E2E8F0"
    heading_font: str = "Calibri"
    body_font: str = "Calibri"


class DocxPlaybookConverter:
    """Conversor de DOCX para Playbook formatado"""

    def __init__(self, brand: BrandIdentity = None):
        self.brand = brand or BrandIdentity()

    def _hex_to_rgb(self, hex_color: str) -> tuple:
        """Converte hex para RGB"""
        hex_color = hex_color.lstrip("#")
        return (
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16)
        )

    def read_docx(self, file_path: Path) -> dict:
        """Le documento DOCX e extrai conteudo"""
        print(f"\n[1/4] Lendo documento: {file_path.name}")

        doc = Document(file_path)

        elements = []
        images = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detectar tipo de paragrafo
            style_name = para.style.name if para.style else "Normal"

            if style_name.startswith("Heading") or text.isupper():
                level = 1
                if "Heading 2" in style_name or len(text) < 50:
                    level = 2
                elif "Heading 3" in style_name:
                    level = 3

                elements.append({
                    "type": "heading",
                    "level": level,
                    "text": text
                })
            else:
                # Detectar se e lista
                is_bullet = text.startswith(("-", "*", "•"))
                is_numbered = re.match(r"^\d+[.\)]\s", text)

                elements.append({
                    "type": "paragraph",
                    "text": text,
                    "is_bullet": is_bullet,
                    "is_numbered": is_numbered
                })

        # Extrair tabelas
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)

            if table_data:
                elements.append({
                    "type": "table",
                    "data": table_data
                })

        # Contar imagens
        image_count = len(doc.part.rels.values())

        return {
            "title": file_path.stem,
            "elements": elements,
            "image_count": image_count
        }

    def rewrite_text(self, text: str) -> str:
        """Reescreve texto em linguagem mais clara"""
        # Regras simples de reescrita
        replacements = {
            "devera": "deve",
            "deverao": "devem",
            "devera ser": "e",
            "deverao ser": "sao",
            "realizar": "fazer",
            "executar": "fazer",
            "inicializar": "abrir",
            "finalizar": "fechar",
            "processo de": "",
            "procedimento de": "",
        }

        result = text
        for old, new in replacements.items():
            result = re.sub(r'\b' + old + r'\b', new, result, flags=re.IGNORECASE)

        # Capitalizar primeira letra
        if result:
            result = result[0].upper() + result[1:]

        return result

    def create_playbook(self, content: dict, output_path: Path):
        """Cria playbook DOCX formatado"""
        print(f"[2/4] Criando playbook: {output_path.name}")

        doc = Document()

        # Configurar estilos
        self._setup_styles(doc)

        # Adicionar titulo principal
        title = doc.add_heading(content["title"], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adicionar linha decorativa
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = separator.add_run("━" * 40)
        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.secondary_color))
        run.font.size = Pt(8)

        # Adicionar metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(f"Playbook Lector | Gerado em {datetime.now().strftime('%d/%m/%Y')}")
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_color))
        meta_run.italic = True

        # Espaco
        doc.add_paragraph()

        # Adicionar caixa de informacao
        info_box = doc.add_paragraph()
        info_box_run = info_box.add_run("🎯 SOBRE ESTE PLAYBOOK")
        info_box_run.bold = True
        info_box_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_color))
        info_box_run.font.size = Pt(12)

        info_desc = doc.add_paragraph()
        info_desc_run = info_desc.add_run(
            f"Este guia foi criado a partir do documento '{content['title']}'. "
            "Siga as instrucoes passo a passo para obter os melhores resultados."
        )
        info_desc_run.font.size = Pt(10)
        info_desc.paragraph_format.left_indent = Inches(0.2)

        doc.add_paragraph()

        # Processar elementos
        step_counter = 0

        for element in content["elements"]:
            if element["type"] == "heading":
                level = element.get("level", 1)
                text = self.rewrite_text(element["text"])

                if level == 1:
                    heading = doc.add_heading(text, level=1)
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_color))
                        run.font.size = Pt(24)
                elif level == 2:
                    heading = doc.add_heading(text, level=2)
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_color))
                        run.font.size = Pt(18)
                else:
                    heading = doc.add_heading(text, level=3)
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.secondary_color))
                        run.font.size = Pt(14)

            elif element["type"] == "paragraph":
                text = self.rewrite_text(element["text"])

                # Detectar se e passo de procedimento
                is_step = self._is_procedure_step(text)

                if is_step:
                    step_counter += 1
                    self._add_step(doc, step_counter, text)
                else:
                    para = doc.add_paragraph(text)
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_after = Pt(8)

            elif element["type"] == "table":
                self._add_table(doc, element["data"])

        # Adicionar footer
        doc.add_page_break()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("━" * 30)
        footer_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.border_color))

        footer_text = doc.add_paragraph()
        footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text_run = footer_text.add_run("Lector Tecnologia | www.lector.com.br")
        footer_text_run.font.size = Pt(9)
        footer_text_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_color))

        # Salvar documento
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        print(f"[3/4] Playbook salvo: {output_path}")

    def _setup_styles(self, doc):
        """Configura estilos do documento"""
        # Estilo Normal
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.brand.body_font
        font.size = Pt(11)
        font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_color))

        # Configurar margens
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    def _is_procedure_step(self, text: str) -> bool:
        """Detecta se o texto e um passo de procedimento"""
        indicators = [
            r"^\d+[.\)]\s",
            r"\b(clique|selecione|digite|preencha|envie|acesse|abra|faca)\b",
            r"\b(proximo passo|em seguida|depois)\b",
        ]

        for pattern in indicators:
            if re.search(pattern, text, re.IGNORECASE):
                return True

        return False

    def _add_step(self, doc, number: int, text: str):
        """Adiciona um passo formatado"""
        # Numero do passo
        step_para = doc.add_paragraph()
        step_run = step_para.add_run(f"PASSO {number:02d}")
        step_run.bold = True
        step_run.font.size = Pt(10)
        step_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.secondary_color))

        # Descricao
        desc_para = doc.add_paragraph()
        desc_run = desc_para.add_run(text)
        desc_run.font.size = Pt(11)
        desc_para.paragraph_format.left_indent = Inches(0.3)
        desc_para.paragraph_format.space_after = Pt(12)

    def _add_table(self, doc, data: list):
        """Adiciona tabela formatada"""
        if not data or len(data) < 1:
            return

        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = "Table Grid"

        for i, row in enumerate(data):
            for j, cell_text in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = cell_text

                # Estilizar header
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)


def main():
    """Funcao principal"""
    print("=" * 60)
    print("  LECTOR PLAYBOOK CONVERTER")
    print("  Conversor Simplificado de DOCX")
    print("=" * 60)

    # Configuracoes
    input_dir = Path(r"C:\Users\Lector\Desktop\Documentações")
    output_dir = Path(r"C:\Users\Lector\lector-playbook-converter\output")

    # Verificar diretorio de entrada
    if not input_dir.exists():
        print(f"\nErro: Pasta nao encontrada: {input_dir}")
        return 1

    # Encontrar arquivos
    docx_files = list(input_dir.glob("*.docx"))

    if not docx_files:
        print(f"\nNenhum arquivo .docx encontrado em {input_dir}")
        return 1

    print(f"\nEncontrados {len(docx_files)} arquivo(s):")
    for f in docx_files:
        print(f"  - {f.name}")

    # Criar conversor
    brand = BrandIdentity()
    converter = DocxPlaybookConverter(brand)

    # Processar cada arquivo
    print("\n" + "-" * 60)
    results = []

    for docx_file in docx_files:
        try:
            print(f"\nProcessando: {docx_file.name}")
            print("-" * 40)

            # Ler documento
            content = converter.read_docx(docx_file)

            # Criar nome de saida
            output_file = output_dir / f"PLAYBOOK_{docx_file.stem}.docx"

            # Criar playbook
            converter.create_playbook(content, output_file)

            results.append((docx_file.name, output_file.name, "Sucesso"))
            print(f"[4/4] ✓ Concluido!")

        except Exception as e:
            print(f"\n[ERRO] Falha ao processar {docx_file.name}: {e}")
            results.append((docx_file.name, None, f"Erro: {e}"))

    # Resumo
    print("\n" + "=" * 60)
    print("  RESUMO")
    print("=" * 60)

    for input_name, output_name, status in results:
        status_icon = "✓" if status == "Sucesso" else "✗"
        print(f"  {status_icon} {input_name}")
        if output_name:
            print(f"    → {output_name}")

    print(f"\n  Playbooks salvos em: {output_dir}")
    print("=" * 60 + "\n")

    return 0


if __name__ == "__main__":
    sys.exit(main())
