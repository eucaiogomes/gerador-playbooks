#!/usr/bin/env python3
"""
Lector Playbook Converter - Conversão em Massa (v2)
Converte todos os DOCX de uma pasta em playbooks organizados com 3 versões:
- HTML (interativo)
- PDF (imprimível)
- Word (.docx)

Estrutura de saída:
output/
├── index.html
├── Playbook-1/
│   ├── playbook.html
│   ├── playbook.pdf
│   ├── playbook.docx
│   └── images/
├── Playbook-2/
│   └── ...

Uso:
  python batch_convert_v2.py
  python batch_convert_v2.py --input "C:/caminho/docs" --output "C:/saida"
"""

import sys
import importlib.util
from pathlib import Path
from datetime import datetime
import shutil

import click
from rich.console import Console
from rich.table import Table
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, MofNCompleteColumn
from rich.panel import Panel
from rich import box

ROOT    = Path(__file__).parent
SRC     = ROOT / "src"
TOOLS   = SRC / "tools"
console = Console()


# ── Carregamento direto dos módulos (evita imports relativos) ────────────────
def _load_mod(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    mod  = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod

_html_mod  = _load_mod("html_writer", TOOLS / "html_writer.py")
_docx_mod  = _load_mod("docx_reader", TOOLS / "docx_reader.py")
HTMLWriter = _html_mod.HTMLWriter
DocxReader = _docx_mod.DocxReader

DEFAULT_INPUT  = Path(r"C:\Users\Lector\Desktop\Documentações")
DEFAULT_OUTPUT = ROOT / "output"
DEFAULT_LOGO   = Path(r"C:\Users\Lector\Downloads\logo\logo-lector.svg")


def slugify(text: str) -> str:
    """Converte texto em nome de pasta seguro"""
    import re
    # Remover extensão .docx
    text = text.replace('.docx', '').replace('.doc', '')
    # Converter caracteres especiais
    text = re.sub(r'[^\w\s-]', '', text)
    # Substituir espaços por hífen
    text = re.sub(r'[\s_]+', '-', text)
    # Remover hífens duplicados
    text = re.sub(r'-+', '-', text)
    return text.strip('-').lower()[:50] or "playbook"


def load_logo(logo_path: Path) -> str:
    """Carrega o SVG da logo"""
    if logo_path.exists():
        return logo_path.read_text(encoding='utf-8')
    console.print(f"[yellow]Aviso: logo não encontrada em {logo_path}[/yellow]")
    return ""


def generate_word_document(doc_content: Dict[str, Any], output_path: Path, title: str):
    """Gera um documento Word (.docx) formatado"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import io

    doc = Document()

    # Configurar estilos
    styles = doc.styles

    # Estilo para título principal
    style_title = styles.add_style('LectorTitle', WD_STYLE_TYPE.PARAGRAPH)
    style_title.font.name = 'Calibri'
    style_title.font.size = Pt(20)
    style_title.font.bold = True
    style_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)  # Navy

    # Estilo para headings
    for i in range(1, 4):
        style_h = styles.add_style(f'LectorHeading{i}', WD_STYLE_TYPE.PARAGRAPH)
        style_h.font.name = 'Calibri'
        style_h.font.size = Pt(16 - i)
        style_h.font.bold = True
        style_h.font.color.rgb = RGBColor(0xEF, 0x63, 0x15) if i == 1 else RGBColor(0x1E, 0x3A, 0x5F)

    # Adicionar título
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adicionar subtítulo
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"Playbook Lector Tecnologia • {datetime.now().strftime('%d/%m/%Y')}")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Espaço

    # Mapear imagens
    images_map = {}
    for img in doc_content.get('images', []):
        try:
            data = img.image_data if hasattr(img, 'image_data') else img.get('image_data', b'')
            h = img.hash if hasattr(img, 'hash') else img.get('hash', '')
            if data and h:
                images_map[h] = data
        except Exception:
            pass

    # Processar elementos
    for el in doc_content.get('elements', []):
        t = el.type if hasattr(el, 'type') else el.get('type', 'paragraph')
        c = el.content if hasattr(el, 'content') else el.get('content', '')
        lv = el.level if hasattr(el, 'level') else el.get('level', 1)

        if t == 'heading':
            # Adicionar heading
            p = doc.add_heading(c, level=min(lv, 3))
            p.style.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

        elif t == 'subheading':
            p = doc.add_heading(c, level=min(lv, 4))
            p.style.font.color.rgb = RGBColor(0x2A, 0x4F, 0x80)

        elif t == 'paragraph':
            p = doc.add_paragraph(c)
            p.style.font.size = Pt(11)
            p.style.font.name = 'Calibri'

        elif t == 'image':
            # Adicionar imagem
            img_hash = c.get('hash', '') if isinstance(c, dict) else ''
            if img_hash in images_map:
                try:
                    img_stream = io.BytesIO(images_map[img_hash])
                    doc.add_picture(img_stream, width=Inches(5.5))
                except Exception:
                    pass

        elif t == 'list':
            items = c if isinstance(c, list) else []
            is_ordered = el.formatting.get('ordered', False) if hasattr(el, 'formatting') else False

            for item in items:
                text = item.get('text', '') if isinstance(item, dict) else str(item)
                if is_ordered:
                    p = doc.add_paragraph(text, style='List Number')
                else:
                    p = doc.add_paragraph(text, style='List Bullet')

        elif t == 'table':
            rows = c if isinstance(c, list) else []
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]) if rows else 1)
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    for j, cell_text in enumerate(row):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = str(cell_text)

    # Salvar
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def generate_pdf_from_html(html_content: str, output_path: Path, title: str):
    """Gera PDF a partir do HTML usando weasyprint ou alternativa"""
    try:
        from weasyprint import HTML, CSS
        from weasyprint.text import font_configuration

        # CSS adicional para PDF
        pdf_css = """
        @page {
            size: A4;
            margin: 2cm;
            @top-center {
                content: "Lector Tecnologia";
                font-size: 9pt;
                color: #718096;
            }
            @bottom-center {
                content: counter(page);
                font-size: 9pt;
            }
        }
        body { font-family: 'Segoe UI', Calibri, Arial, sans-serif; }
        """

        html_obj = HTML(string=html_content)
        html_obj.write_pdf(output_path, stylesheets=[CSS(string=pdf_css)])
        return True
    except ImportError:
        console.print("[yellow]weasyprint não instalado. Pulando geração de PDF.[/yellow]")
        return False
    except Exception as e:
        console.print(f"[red]Erro ao gerar PDF: {e}[/red]")
        return False


def convert_one(docx_path: Path, output_base_dir: Path, logo_svg: str) -> dict:
    """Converte um DOCX para HTML, PDF e Word, organizando em pasta separada"""
    result = {
        "file": docx_path.name,
        "status": "error",
        "folder": None,
        "html_path": None,
        "pdf_path": None,
        "docx_path": None,
        "sections": 0,
        "error": ""
    }

    try:
        # Criar pasta para este playbook
        folder_name = slugify(docx_path.stem)
        playbook_dir = output_base_dir / folder_name
        playbook_dir.mkdir(parents=True, exist_ok=True)

        # Criar subpasta para imagens
        images_dir = playbook_dir / "images"
        images_dir.mkdir(exist_ok=True)

        result["folder"] = str(playbook_dir)

        # Ler documento
        reader = DocxReader(docx_path)
        doc_content = reader.read()

        # Salvar imagens na pasta do playbook
        for img in doc_content.get('images', []):
            try:
                img_data = img.image_data if hasattr(img, 'image_data') else img.get('image_data', b'')
                filename = img.filename if hasattr(img, 'filename') else img.get('filename', 'image.png')
                img_path = images_dir / filename
                img_path.write_bytes(img_data)
            except Exception:
                pass

        # Gerar HTML
        html_name = "playbook.html"
        html_path = playbook_dir / html_name

        writer = HTMLWriter(
            output_path=html_path,
            title=docx_path.stem,
            logo_svg=logo_svg,
        )
        writer.build_from_doc_content(doc_content)

        # Salvar HTML
        html_content = writer._render_full_html()
        html_path.write_text(html_content, encoding='utf-8')
        result["html_path"] = str(html_path)
        result["sections"] = len(writer._sections)

        # Gerar PDF
        pdf_path = playbook_dir / "playbook.pdf"
        if generate_pdf_from_html(html_content, pdf_path, docx_path.stem):
            result["pdf_path"] = str(pdf_path)

        # Gerar Word
        docx_output_path = playbook_dir / "playbook.docx"
        try:
            generate_word_document(doc_content, docx_output_path, docx_path.stem)
            result["docx_path"] = str(docx_output_path)
        except Exception as e:
            console.print(f"[yellow]Erro ao gerar Word para {docx_path.name}: {e}[/yellow]")

        result["status"] = "success"

    except Exception as e:
        result["error"] = str(e)
        import traceback
        console.print(f"[red]Erro detalhado:[/red] {traceback.format_exc()}")

    return result


def generate_index(output_dir: Path, results: list, logo_svg: str) -> Path:
    """Gera um index.html com cards para todos os playbooks e links para as 3 versões"""
    ok = [r for r in results if r["status"] == "success"]
    now = datetime.now().strftime("%d/%m/%Y %H:%M")

    # Logo para o index
    if logo_svg:
        import re
        logo_html = re.sub(r'<svg ', '<svg style="height:38px;width:auto" ', logo_svg.strip(), count=1)
    else:
        logo_html = '<span style="font-size:20px;font-weight:800;color:#fff"><span style="color:#EF6315">iT</span> Lector</span>'

    cards = ""
    for r in ok:
        name = r["file"].replace(".docx", "").replace(".doc", "")
        folder_name = slugify(name)

        # Verificar quais arquivos existem
        has_html = r.get("html_path") and Path(r["html_path"]).exists()
        has_pdf = r.get("pdf_path") and Path(r["pdf_path"]).exists()
        has_docx = r.get("docx_path") and Path(r["docx_path"]).exists()

        # Botões de download
        buttons = []
        if has_html:
            buttons.append(f'<a href="{folder_name}/playbook.html" class="btn btn-primary" target="_blank">👁️ Visualizar</a>')
        if has_pdf:
            buttons.append(f'<a href="{folder_name}/playbook.pdf" class="btn btn-secondary" download>📄 PDF</a>')
        if has_docx:
            buttons.append(f'<a href="{folder_name}/playbook.docx" class="btn btn-secondary" download>📝 Word</a>')

        buttons_html = ''.join(buttons) if buttons else '<span class="no-files">Nenhum arquivo gerado</span>'

        cards += f"""
        <div class="card">
          <div class="card-icon">📋</div>
          <div class="card-body">
            <h3 class="card-title">{name}</h3>
            <div class="card-meta">
              <span>📑 {r.get("sections", 0)} seções</span>
            </div>
            <div class="card-actions">
              {buttons_html}
            </div>
          </div>
        </div>"""

    errors = [r for r in results if r["status"] == "error"]
    err_section = ""
    if errors:
        err_rows = ""
        for r in errors:
            err_rows += f'<tr><td>{r["file"]}</td><td class="err">{r.get("error", "")[:80]}</td></tr>'
        err_section = f"""
        <div class="err-block">
          <h2>Arquivos com Erro ({len(errors)})</h2>
          <table class="err-tbl"><thead><tr><th>Arquivo</th><th>Erro</th></tr></thead>
          <tbody>{err_rows}</tbody></table>
        </div>"""

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Lector Playbooks - Central de Documentação</title>
<style>
:root {{
  --navy:#00347E; --navy-dark:#001C66; --orange:#EF6315; --orange2:#FF930F;
  --white:#fff; --gray-50:#F7FAFC; --gray-100:#EDF2F7; --gray-200:#E2E8F0;
  --gray-600:#718096; --gray-700:#4A5568; --success:#38A169;
  --r:10px; --sh:0 2px 8px rgba(0,0,0,.1);
}}
*,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'Segoe UI',Calibri,Arial,sans-serif;background:var(--gray-50);color:var(--gray-700);min-height:100vh}}

/* Header */
.hdr{{background:var(--navy-dark);padding:0 32px;height:64px;display:flex;align-items:center;gap:14px;box-shadow:0 2px 8px rgba(0,0,0,.3)}}
.hdr svg{{height:34px;width:auto}}
.hdr-div{{width:1px;height:26px;background:rgba(255,255,255,.15)}}
.hdr-title{{font-size:14px;font-weight:500;color:rgba(255,255,255,.75);flex:1}}
.hdr-date{{font-size:12px;color:rgba(255,255,255,.4)}}

/* Hero */
.hero{{background:linear-gradient(135deg,var(--navy-dark) 0%,var(--navy) 100%);padding:48px 32px 40px;text-align:center;color:var(--white)}}
.hero h1{{font-size:32px;font-weight:700;margin-bottom:8px}}
.hero p{{font-size:15px;color:rgba(255,255,255,.7);margin-bottom:24px}}
.hero-stats{{display:inline-flex;gap:32px;background:rgba(255,255,255,.08);border:1px solid rgba(255,255,255,.12);border-radius:var(--r);padding:14px 32px}}
.hstat{{text-align:center}}
.hstat-val{{font-size:26px;font-weight:700;color:var(--orange);display:block}}
.hstat-lbl{{font-size:11px;color:rgba(255,255,255,.55);text-transform:uppercase;letter-spacing:.5px}}

/* Content */
.content{{max-width:1100px;margin:0 auto;padding:40px 24px}}
.section-title{{font-size:13px;font-weight:700;color:var(--gray-700);text-transform:uppercase;letter-spacing:.8px;margin-bottom:16px}}

/* Cards Grid */
.cards{{display:grid;grid-template-columns:repeat(auto-fill,minmax(320px,1fr));gap:16px;margin-bottom:40px}}
.card{{display:flex;flex-direction:column;background:var(--white);border-radius:var(--r);padding:20px;box-shadow:var(--sh);border:1px solid var(--gray-200);transition:all .2s ease}}
.card:hover{{box-shadow:0 6px 20px rgba(0,0,0,.12);transform:translateY(-2px);border-color:var(--orange)}}
.card-icon{{font-size:28px;margin-bottom:12px}}
.card-body{{flex:1}}
.card-title{{font-size:16px;font-weight:600;color:var(--navy);margin-bottom:8px;line-height:1.3}}
.card-meta{{display:flex;gap:10px;font-size:12px;color:var(--gray-600);margin-bottom:14px}}
.card-actions{{display:flex;gap:8px;flex-wrap:wrap}}

/* Buttons */
.btn{{display:inline-flex;align-items:center;gap:5px;padding:8px 14px;border-radius:6px;font-size:13px;font-weight:500;text-decoration:none;transition:all .2s ease;border:1px solid transparent;cursor:pointer}}
.btn-primary{{background:var(--orange);color:var(--white);border-color:var(--orange)}}
.btn-primary:hover{{background:var(--orange2);border-color:var(--orange2)}}
.btn-secondary{{background:var(--gray-100);color:var(--gray-700);border-color:var(--gray-200)}}
.btn-secondary:hover{{background:var(--gray-200);border-color:var(--gray-300)}}

.no-files{{font-size:12px;color:var(--gray-600);font-style:italic}}

/* Errors */
.err-block{{margin-top:32px}}
.err-block h2{{font-size:16px;font-weight:600;color:#E53E3E;margin-bottom:12px}}
.err-tbl{{width:100%;border-collapse:collapse;font-size:13px;background:var(--white);border-radius:var(--r);overflow:hidden;box-shadow:var(--sh)}}
.err-tbl th{{background:#FFF5F5;color:#E53E3E;padding:10px 14px;text-align:left;font-size:11px;text-transform:uppercase}}
.err-tbl td{{padding:9px 14px;border-bottom:1px solid var(--gray-200)}}
.err{{color:#E53E3E;font-size:12px}}

/* Footer */
.footer{{text-align:center;padding:24px;font-size:12px;color:#A0AEC0;border-top:1px solid var(--gray-200)}}

/* Responsive */
@media(max-width:600px){{
  .hero{{padding:32px 16px 28px}}
  .hero-stats{{gap:20px;padding:12px 20px}}
  .content{{padding:24px 14px}}
  .cards{{grid-template-columns:1fr}}
  .hdr{{padding:0 16px}}
}}
</style>
</head>
<body>
<header class="hdr">
  {logo_html}
  <div class="hdr-div"></div>
  <span class="hdr-title">Central de Playbooks</span>
  <span class="hdr-date">{now}</span>
</header>

<div class="hero">
  <h1>Central de Playbooks</h1>
  <p>Documentação interativa da equipe Lector Tecnologia</p>
  <div class="hero-stats">
    <div class="hstat"><span class="hstat-val">{len(ok)}</span><span class="hstat-lbl">Playbooks</span></div>
    <div class="hstat"><span class="hstat-val">{sum(r.get('sections',0) for r in ok)}</span><span class="hstat-lbl">Seções</span></div>
  </div>
</div>

<div class="content">
  <div class="section-title">Playbooks Disponíveis ({len(ok)})</div>
  <div class="cards">{cards}</div>
  {err_section}
</div>

<footer class="footer">© Lector Tecnologia · Gerado em {now}</footer>
</body>
</html>"""

    index_path = output_dir / "index.html"
    index_path.write_text(html, encoding='utf-8')
    return index_path


# ── CLI ──────────────────────────────────────────────────────────────────────

@click.command()
@click.option('--input',  '-i', 'input_dir',  default=str(DEFAULT_INPUT),
              type=click.Path(), show_default=True, help='Pasta com arquivos .docx')
@click.option('--output', '-o', 'output_dir', default=str(DEFAULT_OUTPUT),
              type=click.Path(), show_default=True, help='Pasta de saída')
@click.option('--logo',   '-l', 'logo_path',  default=str(DEFAULT_LOGO),
              type=click.Path(), show_default=True, help='Arquivo SVG da logo Lector')
@click.option('--pattern', '-p', default='**/*.docx',
              show_default=True, help='Glob pattern para encontrar arquivos')
def main(input_dir, output_dir, logo_path, pattern):
    """
    Converte todos os arquivos DOCX em playbooks organizados com 3 versões:
    HTML (interativo), PDF (imprimível) e Word (.docx).
    Cria uma pasta separada para cada playbook.
    """
    input_path  = Path(input_dir)
    output_path = Path(output_dir)
    logo_file   = Path(logo_path)

    console.print(Panel.fit(
        "[bold cyan]Lector Playbook Converter[/bold cyan] — [dim]Versão 2 - Organizado[/dim]",
        border_style="cyan"
    ))

    if not input_path.exists():
        console.print(f"[red]Pasta de entrada não encontrada: {input_path}[/red]")
        raise SystemExit(1)

    output_path.mkdir(parents=True, exist_ok=True)
    logo_svg = load_logo(logo_file)

    docx_files = sorted(input_path.glob(pattern))
    if not docx_files:
        console.print(f"[yellow]Nenhum .docx encontrado em: {input_path}[/yellow]")
        raise SystemExit(0)

    console.print(f"[blue]Encontrados:[/blue] {len(docx_files)} arquivo(s)")
    console.print(f"[dim]Saída será organizada em pastas separadas em:[/dim] {output_path.resolve()}\n")

    results = []
    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}", table_column=None),
        BarColumn(bar_width=30),
        MofNCompleteColumn(),
        console=console,
    ) as prog:
        task = prog.add_task("Convertendo...", total=len(docx_files))
        for f in docx_files:
            prog.update(task, description=f"[cyan]{f.name[:40]}[/cyan]")
            r = convert_one(f, output_path, logo_svg)
            results.append(r)
            prog.advance(task)

    # Index page
    idx = generate_index(output_path, results, logo_svg)

    # Summary table
    table = Table(title="Resultado da Conversão", box=box.ROUNDED)
    table.add_column("Playbook", style="cyan", max_width=35)
    table.add_column("Status", justify="center", width=8)
    table.add_column("Pasta", width=20)
    table.add_column("HTML", justify="center", width=6)
    table.add_column("PDF", justify="center", width=6)
    table.add_column("Word", justify="center", width=6)

    for r in results:
        st = "[green]✓[/green]" if r["status"] == "success" else "[red]✗[/red]"
        folder = Path(r["folder"]).name if r["folder"] else "-"
        has_html = "✓" if r.get("html_path") else "-"
        has_pdf = "✓" if r.get("pdf_path") else "-"
        has_docx = "✓" if r.get("docx_path") else "-"
        table.add_row(r["file"][:35], st, folder, has_html, has_pdf, has_docx)

    console.print()
    console.print(table)

    ok = sum(1 for r in results if r["status"] == "success")
    err = len(results) - ok
    console.print(f"\n[green]{ok} playbooks gerados com sucesso[/green]", end="")
    if err:
        console.print(f"  [red]{err} com erro[/red]", end="")
    console.print(f"\n[dim]Local:[/dim] {output_path.resolve()}")
    console.print(f"[bold]Índice:[/bold] {idx.resolve()}\n")

    # Mostrar estrutura
    console.print("[dim]Estrutura gerada:[/dim]")
    for r in results:
        if r["status"] == "success" and r["folder"]:
            folder = Path(r["folder"])
            console.print(f"  📁 {folder.name}/")
            if r.get("html_path"):
                console.print(f"     📄 playbook.html")
            if r.get("pdf_path"):
                console.print(f"     📄 playbook.pdf")
            if r.get("docx_path"):
                console.print(f"     📄 playbook.docx")
            console.print(f"     📁 images/")


if __name__ == "__main__":
    main()
