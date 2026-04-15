#!/usr/bin/env python3
"""
Lector Playbook Converter - Versão Visual com Imagens
Extrai imagens dos documentos originais e aplica design no estilo Lector
"""

import os
import sys
import re
import io
import hashlib
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass
from typing import List, Dict, Tuple

sys.path.insert(0, str(Path(__file__).parent / "src"))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    from PIL import Image as PILImage
except ImportError as e:
    print(f"Erro ao importar: {e}")
    os.system("pip install python-docx pillow")
    sys.exit(1)


@dataclass
class BrandIdentity:
    """Identidade visual Lector extraída do site"""
    primary_dark: str = "#1E3A5F"
    primary_bright: str = "#2563EB"
    accent_coral: str = "#D97757"
    text_dark: str = "#2D3748"
    text_secondary: str = "#4A5568"
    light_bg: str = "#F7FAFC"
    border: str = "#E2E8F0"
    white: str = "#FFFFFF"


class ImageExtractor:
    """Extrai imagens de arquivos DOCX"""

    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.images_dir = output_dir / "images"
        self.images_dir.mkdir(parents=True, exist_ok=True)

    def extract_images(self, doc_path: Path) -> List[Dict]:
        """Extrai todas as imagens de um DOCX"""
        print(f"  Extraindo imagens de: {doc_path.name}")

        doc = Document(doc_path)
        images = []
        image_count = 0

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image = rel.target_part
                    image_data = image.blob
                    content_type = image.content_type

                    # Determinar extensão
                    if "png" in content_type:
                        ext = "png"
                    elif "jpeg" in content_type or "jpg" in content_type:
                        ext = "jpg"
                    elif "gif" in content_type:
                        ext = "gif"
                    else:
                        ext = "png"

                    # Criar hash único
                    image_hash = hashlib.md5(image_data).hexdigest()[:8]
                    filename = f"img_{doc_path.stem}_{image_hash}.{ext}"
                    filepath = self.images_dir / filename

                    # Salvar imagem
                    with open(filepath, "wb") as f:
                        f.write(image_data)

                    # Obter dimensões
                    try:
                        img = PILImage.open(io.BytesIO(image_data))
                        width, height = img.size
                    except:
                        width, height = None, None

                    images.append({
                        "path": filepath,
                        "filename": filename,
                        "data": image_data,
                        "width": width,
                        "height": height,
                        "extension": ext
                    })
                    image_count += 1

                except Exception as e:
                    print(f"    Aviso: Não foi possível extrair uma imagem: {e}")

        print(f"  {image_count} imagem(s) extraída(s)")
        return images


class VisualPlaybookConverter:
    """Conversor com design visual no estilo Lector"""

    def __init__(self, brand: BrandIdentity = None):
        self.brand = brand or BrandIdentity()
        self.image_extractor = None

    def _hex_to_rgb(self, hex_color: str) -> tuple:
        """Converte hex para RGB"""
        hex_color = hex_color.lstrip("#")
        return (
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16)
        )

    def read_docx(self, file_path: Path) -> dict:
        """Lê documento DOCX e extrai conteúdo estruturado"""
        print(f"\n[1/5] Analisando documento: {file_path.name}")

        doc = Document(file_path)
        elements = []

        # Extrair parágrafos
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else "Normal"

            # Detectar títulos
            if style_name.startswith("Heading"):
                level = int(style_name.replace("Heading ", "")) if "Heading " in style_name else 2
                elements.append({
                    "type": "heading",
                    "level": level,
                    "text": text
                })
            elif text.isupper() and len(text) < 100:
                elements.append({
                    "type": "heading",
                    "level": 2,
                    "text": text
                })
            else:
                # Detectar listas
                if text.startswith(("-", "*", "•")):
                    elements.append({
                        "type": "bullet",
                        "text": text[1:].strip()
                    })
                elif re.match(r"^\d+[.\)]\s", text):
                    elements.append({
                        "type": "numbered",
                        "text": re.sub(r"^\d+[.\)]\s", "", text)
                    })
                else:
                    elements.append({
                        "type": "paragraph",
                        "text": text
                    })

        # Extrair tabelas
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)

            if table_data:
                elements.append({
                    "type": "table",
                    "data": table_data
                })

        return {
            "title": file_path.stem,
            "elements": elements
        }

    def rewrite_text(self, text: str) -> str:
        """Reescreve texto em linguagem mais clara"""
        if not text:
            return text

        replacements = {
            "devera": "deve",
            "deverao": "devem",
            "devera ser": "é",
            "deverao ser": "são",
            "realizar": "fazer",
            "executar": "fazer",
            "inicializar": "abrir",
            "finalizar": "fechar",
            "processo de": "",
            "procedimento de": "",
            "atraves de": "por meio de",
            "no que diz respeito": "sobre",
            "no intuito de": "para",
            "visando": "para",
            "com o objetivo de": "para",
        }

        result = text
        for old, new in replacements.items():
            result = re.sub(r'\b' + old + r'\b', new, result, flags=re.IGNORECASE)

        # Capitalizar primeira letra
        if result:
            result = result[0].upper() + result[1:]

        return result

    def create_visual_playbook(self, content: dict, images: List[Dict], output_path: Path):
        """Cria playbook com design visual Lector"""
        print(f"[2/5] Criando playbook visual: {output_path.name}")

        doc = Document()

        # Configurar estilos e margens
        self._setup_document_styles(doc)

        # ========== CAPA / HEADER ==========
        self._add_header_section(doc, content["title"])

        # ========== SOBRE O PLAYBOOK ==========
        self._add_about_section(doc)

        # ========== CONTEÚDO PRINCIPAL ==========
        print(f"[3/5] Formatando conteúdo...")
        self._add_content_sections(doc, content["elements"])

        # ========== IMAGENS DE REFERÊNCIA ==========
        if images:
            print(f"[4/5] Inserindo {len(images)} imagem(s)...")
            self._add_images_section(doc, images)

        # ========== FOOTER ==========
        self._add_footer_section(doc)

        # Salvar
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        print(f"[5/5] Playbook salvo!")

    def _setup_document_styles(self, doc):
        """Configura estilos do documento"""
        # Margens
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Estilo Normal
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_dark))

        # Estilo Título
        title_style = doc.styles["Title"]
        title_style.font.name = "Calibri"
        title_style.font.size = Pt(32)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_dark))

        # Estilos de Heading
        for i in range(1, 4):
            h_style = doc.styles[f"Heading {i}"]
            h_style.font.name = "Calibri"
            h_style.font.bold = True

    def _add_header_section(self, doc, title: str):
        """Adiciona capa visual estilo Lector"""
        # Espaço no topo
        doc.add_paragraph()

        # Box de destaque azul escuro
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run("PLAYBOOK")
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_bright))

        # Título principal
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title.upper())
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_dark))

        # Linha decorativa
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run("▬" * 30)
        line_run.font.size = Pt(8)
        line_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_bright))

        # Subtítulo
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run("Guia Prático | Lector Tecnologia")
        sub_run.font.size = Pt(11)
        sub_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_secondary))
        sub_run.italic = True

        # Espaço
        doc.add_paragraph()
        doc.add_paragraph()

    def _add_about_section(self, doc):
        """Adiciona caixa 'Sobre este Playbook' estilo framework"""
        # Título da seção
        about_title = doc.add_paragraph()
        about_title_run = about_title.add_run("🎯 SOBRE ESTE PLAYBOOK")
        about_title_run.font.size = Pt(14)
        about_title_run.font.bold = True
        about_title_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_dark))

        # Linha decorativa
        about_line = doc.add_paragraph()
        about_line_run = about_line.add_run("━" * 50)
        about_line_run.font.size = Pt(6)
        about_line_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_bright))

        # Descrição
        about_text = doc.add_paragraph()
        about_text_run = about_text.add_run(
            "Este guia foi desenvolvido para ser prático, direto e fácil de seguir. "
            "Cada seção apresenta informações estruturadas para facilitar a execução das tarefas. "
            "Siga os passos na ordem apresentada para obter os melhores resultados."
        )
        about_text_run.font.size = Pt(10)
        about_text.paragraph_format.left_indent = Inches(0.2)

        # Espaço
        doc.add_paragraph()

    def _add_content_sections(self, doc, elements: List[Dict]):
        """Adiciona seções de conteúdo formatadas"""
        step_counter = 0
        current_section = None

        for element in elements:
            if element["type"] == "heading":
                level = element.get("level", 1)
                text = self.rewrite_text(element["text"])

                if level == 1:
                    # Título principal de seção - estilo destaque
                    doc.add_paragraph()  # Espaço antes
                    heading = doc.add_heading(text, level=1)
                    for run in heading.runs:
                        run.font.size = Pt(20)
                        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_dark))
                    current_section = text

                elif level == 2:
                    # Subtítulo
                    heading = doc.add_heading(text, level=2)
                    for run in heading.runs:
                        run.font.size = Pt(16)
                        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_dark))

                else:
                    # Tópico
                    heading = doc.add_heading(text, level=3)
                    for run in heading.runs:
                        run.font.size = Pt(13)
                        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_bright))

            elif element["type"] == "bullet":
                text = self.rewrite_text(element["text"])
                para = doc.add_paragraph(text, style="List Bullet")
                para.paragraph_format.left_indent = Inches(0.3)
                para.paragraph_format.space_after = Pt(6)

            elif element["type"] == "numbered":
                text = self.rewrite_text(element["text"])
                para = doc.add_paragraph(text, style="List Number")
                para.paragraph_format.left_indent = Inches(0.3)
                para.paragraph_format.space_after = Pt(6)

            elif element["type"] == "paragraph":
                text = self.rewrite_text(element["text"])

                # Detectar se é passo de procedimento
                if self._is_procedure_step(text):
                    step_counter += 1
                    self._add_step_box(doc, step_counter, text)
                else:
                    para = doc.add_paragraph(text)
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_after = Pt(8)

            elif element["type"] == "table":
                self._add_styled_table(doc, element["data"])

    def _add_step_box(self, doc, number: int, text: str):
        """Adiciona caixa de passo estilo framework"""
        # Espaço antes
        doc.add_paragraph()

        # Box do passo
        step_header = doc.add_paragraph()
        step_header.paragraph_format.left_indent = Inches(0.2)

        # Número do passo em destaque
        num_run = step_header.add_run(f"PASSO {number:02d} ")
        num_run.font.size = Pt(11)
        num_run.font.bold = True
        num_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_bright))

        # Ícone
        icon_run = step_header.add_run("▸ ")
        icon_run.font.size = Pt(11)
        icon_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.accent_coral))

        # Texto do passo
        text_run = step_header.add_run(text)
        text_run.font.size = Pt(11)
        text_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_dark))

        # Espaço depois
        doc.add_paragraph()

    def _add_styled_table(self, doc, data: List[List[str]]):
        """Adiciona tabela com estilo Lector"""
        if not data or len(data) < 1:
            return

        # Criar tabela
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = "Table Grid"

        # Aplicar dados
        for i, row in enumerate(data):
            for j, cell_text in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = self.rewrite_text(cell_text)

                # Estilizar célula
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

                # Header row - fundo azul escuro
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)

                    # Aplicar cor de fundo via XML
                    shading_elm = parse_xml(
                        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                        f'w:fill="{self.brand.primary_dark.lstrip("#")}"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading_elm)

        # Espaço depois da tabela
        doc.add_paragraph()

    def _add_images_section(self, doc, images: List[Dict]):
        """Adiciona seção de imagens de referência"""
        # Quebra de página
        doc.add_page_break()

        # Título da seção
        img_section = doc.add_paragraph()
        img_section_run = img_section.add_run("📸 IMAGENS DE REFERÊNCIA")
        img_section_run.font.size = Pt(16)
        img_section_run.font.bold = True
        img_section_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_dark))

        # Descrição
        img_desc = doc.add_paragraph()
        img_desc_run = img_desc.add_run(
            "As imagens abaixo foram extraídas do documento original e "
            "são apresentadas aqui para complementar as instruções."
        )
        img_desc_run.font.size = Pt(10)
        img_desc_run.font.italic = True
        img_desc.paragraph_format.space_after = Pt(12)

        # Inserir cada imagem
        for i, img in enumerate(images, 1):
            try:
                # Legenda
                caption = doc.add_paragraph()
                caption_run = caption.add_run(f"Figura {i}: {img['filename']}")
                caption_run.font.size = Pt(10)
                caption_run.font.bold = True
                caption_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_secondary))

                # Inserir imagem
                with open(img["path"], "rb") as f:
                    image_stream = io.BytesIO(f.read())

                    # Calcular largura proporcional (máx 6 polegadas)
                    if img["width"] and img["height"]:
                        max_width = Inches(5.5)
                        aspect_ratio = img["height"] / img["width"]

                        # Limitar altura máxima
                        if aspect_ratio > 1.5:  # Imagem alta/portrait
                            max_width = Inches(3.5)
                    else:
                        max_width = Inches(5)

                    doc.add_picture(image_stream, width=max_width)

                # Centralizar imagem
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Espaço
                doc.add_paragraph()

            except Exception as e:
                print(f"    Aviso: Não foi possível inserir imagem {img['filename']}: {e}")

    def _add_footer_section(self, doc):
        """Adiciona rodapé estilizado"""
        doc.add_page_break()

        # Linha decorativa
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line.add_run("━" * 40)
        line_run.font.size = Pt(8)
        line_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.border))

        # Logo/texto Lector
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("LECTOR TECNOLOGIA")
        footer_run.font.size = Pt(12)
        footer_run.font.bold = True
        footer_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_dark))

        # Subtexto
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run("www.lector.com.br | Suporte técnico especializado")
        sub_run.font.size = Pt(9)
        sub_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_secondary))

        # Data
        date = doc.add_paragraph()
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date.add_run(f"Documento gerado em {datetime.now().strftime('%d/%m/%Y')}")
        date_run.font.size = Pt(8)
        date_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_secondary))

    def _is_procedure_step(self, text: str) -> bool:
        """Detecta se texto é um passo de procedimento"""
        indicators = [
            r"\b(clique|selecione|digite|preencha|envie|acesse|abra|faça|clicar|selecionar)\b",
            r"\b(proximo passo|em seguida|depois|posteriormente)\b",
            r"\b(insira|adicione|remova|exclua|atualize)\b",
        ]

        for pattern in indicators:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False


def main():
    """Função principal"""
    print("=" * 65)
    print("  LECTOR PLAYBOOK CONVERTER - VERSÃO VISUAL COM IMAGENS")
    print("  Design no estilo Lector | Framework style")
    print("=" * 65)

    # Configurações
    input_dir = Path(r"C:\Users\Lector\Desktop\Documentações")
    output_dir = Path(r"C:\Users\Lector\lector-playbook-converter\output")

    if not input_dir.exists():
        print(f"\nErro: Pasta não encontrada: {input_dir}")
        return 1

    # Encontrar arquivos
    docx_files = list(input_dir.glob("*.docx"))
    if not docx_files:
        print(f"\nNenhum arquivo .docx encontrado")
        return 1

    print(f"\nEncontrados {len(docx_files)} documento(s):")
    for f in docx_files:
        print(f"  • {f.name}")

    # Criar conversor
    brand = BrandIdentity()
    converter = VisualPlaybookConverter(brand)
    converter.image_extractor = ImageExtractor(output_dir)

    # Processar
    print("\n" + "─" * 65)
    results = []

    for docx_file in docx_files:
        try:
            print(f"\n▶ Processando: {docx_file.name}")
            print("─" * 50)

            # Extrair conteúdo
            content = converter.read_docx(docx_file)

            # Extrair imagens
            images = converter.image_extractor.extract_images(docx_file)

            # Criar playbook
            output_file = output_dir / f"PLAYBOOK_{docx_file.stem}.docx"
            converter.create_visual_playbook(content, images, output_file)

            results.append((docx_file.name, output_file.name, len(images), "Sucesso"))
            print(f"✓ Concluído: {output_file.name}")

        except Exception as e:
            print(f"✗ Erro ao processar {docx_file.name}: {e}")
            import traceback
            traceback.print_exc()
            results.append((docx_file.name, None, 0, f"Erro: {e}"))

    # Resumo
    print("\n" + "=" * 65)
    print("  RESUMO DA CONVERSÃO")
    print("=" * 65)

    for input_name, output_name, img_count, status in results:
        if status == "Sucesso":
            print(f"  ✓ {input_name}")
            print(f"    → {output_name}")
            print(f"    → {img_count} imagem(ns) extraída(s)")
        else:
            print(f"  ✗ {input_name} - {status}")

    print(f"\n  📁 Pasta de saída: {output_dir}")
    print("=" * 65 + "\n")

    return 0


if __name__ == "__main__":
    sys.exit(main())
