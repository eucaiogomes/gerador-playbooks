"""
Ferramenta para extrair conteúdo de arquivos DOCX mantendo estrutura, imagens e formatação
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
import hashlib
import re


@dataclass
class DocumentImage:
    """Representa uma imagem extraída"""
    image_data: bytes
    filename: str
    hash: str
    width: Optional[int] = None
    height: Optional[int] = None


@dataclass
class DocumentElement:
    """Representa um elemento genérico do documento"""
    type: str
    content: Any
    style: Optional[str] = None
    level: Optional[int] = None
    formatting: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ListItem:
    """Representa um item de lista"""
    text: str
    level: int  # Nível de indentação (0, 1, 2...)
    is_ordered: bool  # True = numerado, False = marcador
    numbering_id: Optional[int] = None  # ID da lista para agrupar items
    list_index: int = 0  # Índice dentro da lista


@dataclass
class Heading:
    """Representa um título/cabeçalho"""
    text: str
    level: int  # 1-9 para heading levels
    style: Optional[str] = None


@dataclass
class Paragraph:
    """Representa um parágrafo de texto"""
    text: str
    style: Optional[str] = None
    alignment: Optional[str] = None
    is_bold: bool = False
    is_italic: bool = False


class DocxReader:
    """Lê arquivos DOCX e extrai conteúdo estruturado preservando formatação"""

    def __init__(self, file_path: Path):
        self.file_path = Path(file_path)
        self.doc = Document(self.file_path)
        self.images: List[DocumentImage] = []

    def _get_paragraph_alignment(self, paragraph) -> Optional[str]:
        """Extrai o alinhamento do parágrafo"""
        try:
            if paragraph.alignment:
                alignment_map = {
                    WD_ALIGN_PARAGRAPH.LEFT: 'left',
                    WD_ALIGN_PARAGRAPH.CENTER: 'center',
                    WD_ALIGN_PARAGRAPH.RIGHT: 'right',
                    WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
                }
                return alignment_map.get(paragraph.alignment, 'left')
        except Exception:
            pass
        return 'left'

    def _is_list_item(self, paragraph) -> Tuple[bool, bool, int]:
        """
        Detecta se o parágrafo é um item de lista.
        Retorna: (is_list, is_ordered, level)
        """
        try:
            p = paragraph._p
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                return False, False, 0

            # Verificar numPr (numeração)
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                ilvl = numPr.find(qn('w:ilvl'))
                level = 0
                if ilvl is not None:
                    level = int(ilvl.get(qn('w:val'), 0))

                # Verificar se é ordenado (numId presente)
                numId = numPr.find(qn('w:numId'))
                if numId is not None:
                    return True, True, level

            # Verificar marcadores (bullet)
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                style_val = pStyle.get(qn('w:val'), '')
                if 'List' in style_val or 'Bullet' in style_val:
                    # Extrair nível do estilo (ListBullet1, ListBullet2, etc.)
                    match = re.search(r'(\d+)', style_val)
                    level = int(match.group(1)) - 1 if match else 0
                    return True, False, level

            # Verificar pelo texto (heurística)
            text = paragraph.text.strip()
            if text:
                # Padrões de marcadores comuns (bullet, triangle, circle, square, etc)
                bullet_patterns = [
                    r'^[\u2022\u2023\u25E6\u25AA\u25AB\u2013\u2014\-–]\s+',  # Marcadores unicode
                    r'^\[\s*[xX\s]\s*\]\s+',  # [ ], [x], [X]
                    r'^\(\s*[xX\s]\s*\)\s+',  # ( ), (x), (X)
                ]
                for pattern in bullet_patterns:
                    if re.match(pattern, text):
                        return True, False, 0

                # Padrões de numeração (1., 1), (1), A., a., I., i., etc.)
                numbered_patterns = [
                    r'^\d+[\.\)]\s+',  # 1., 1), 2., 2)
                    r'^\([\d\w]\)\s+',  # (1), (a), (A)
                    r'^[a-zA-Z][\.\)]\s+',  # A., a., B., b.
                    r'^[IVXivx]+[\.\)]\s+',  # I., II., i., ii. (romanos)
                ]
                for pattern in numbered_patterns:
                    if re.match(pattern, text):
                        return True, True, 0

        except Exception as e:
            pass

        return False, False, 0

    def _get_heading_level(self, paragraph) -> int:
        """
        Determina o nível de heading do parágrafo.
        Retorna 0 se não for um heading, ou 1-9 para heading levels.
        """
        style_name = ""
        try:
            if paragraph.style and paragraph.style.name:
                style_name = paragraph.style.name
        except Exception:
            pass

        # Verificar estilos de título do Word
        if style_name.startswith("Heading "):
            try:
                level = int(style_name.split(" ")[1])
                return level if 1 <= level <= 9 else 2
            except (ValueError, IndexError):
                pass

        # Verificar estilos em português
        pt_styles = {
            'Título': 1,
            'Título 1': 1, 'Título 2': 2, 'Título 3': 3,
            'Título 4': 4, 'Título 5': 5,
            'Title': 1,
            'Subtitle': 2,
            'Subtítulo': 2,
        }
        if style_name in pt_styles:
            return pt_styles[style_name]

        # Heurísticas para detectar títulos
        text = paragraph.text.strip()
        if not text:
            return 0

        # Se o texto está todo em maiúsculas e é curto
        if text.isupper() and len(text) < 100:
            # Verificar se é fonte maior
            try:
                if paragraph.runs:
                    font_size = None
                    for run in paragraph.runs:
                        if run.font and run.font.size:
                            font_size = run.font.size.pt
                            break
                    if font_size and font_size >= 14:
                        return 2
            except Exception:
                pass
            return 2  # Assume nível 2 para textos em maiúsculas

        # Verificar por padrões de título no texto
        if len(text) < 80:
            # Verificar se é negrito e maior
            try:
                if paragraph.runs:
                    is_bold = any(run.bold for run in paragraph.runs if run.bold is not None)
                    font_size = None
                    for run in paragraph.runs:
                        if run.font and run.font.size:
                            font_size = run.font.size.pt
                            break
                    if is_bold and font_size and font_size >= 14:
                        return 2
                    if is_bold and font_size and font_size >= 12:
                        return 3
            except Exception:
                pass

        return 0

    def _extract_image_from_run(self, run, image_map: Dict[str, bytes]) -> Optional[DocumentImage]:
        """Extrai imagem de um run se houver"""
        try:
            # Verificar se há drawing no run
            run_elem = run._r
            if run_elem is None:
                return None

            # Procurar por blip (imagem)
            for blip in run_elem.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed and embed in image_map:
                    img_data = image_map[embed]
                    img_hash = hashlib.md5(img_data).hexdigest()[:8]

                    # Determinar extensão
                    ext = "png"
                    try:
                        rel = self.doc.part.rels.get(embed)
                        if rel:
                            ct = rel.target_part.content_type
                            if "jpeg" in ct or "jpg" in ct:
                                ext = "jpg"
                            elif "gif" in ct:
                                ext = "gif"
                            elif "png" in ct:
                                ext = "png"
                    except Exception:
                        pass

                    filename = f"img_{img_hash}.{ext}"
                    return DocumentImage(
                        image_data=img_data,
                        filename=filename,
                        hash=img_hash
                    )
        except Exception as e:
            pass
        return None

    def _extract_paragraph_formatting(self, paragraph) -> Dict[str, Any]:
        """Extrai informações detalhadas de formatação do parágrafo"""
        formatting = {
            'style': paragraph.style.name if paragraph.style else None,
            'alignment': self._get_paragraph_alignment(paragraph),
            'text': paragraph.text,
        }

        # Extrair formatação dos runs
        runs_info = []
        for run in paragraph.runs:
            run_info = {
                'text': run.text,
                'bold': run.bold if run.bold is not None else False,
                'italic': run.italic if run.italic is not None else False,
                'underline': run.underline if run.underline is not None else False,
                'font_name': run.font.name if run.font and run.font.name else None,
                'font_size': run.font.size.pt if run.font and run.font.size else None,
                'color': None,
            }

            # Tentar extrair cor
            try:
                if run.font and run.font.color and run.font.color.rgb:
                    run_info['color'] = str(run.font.color.rgb)
            except Exception:
                pass

            runs_info.append(run_info)

        formatting['runs'] = runs_info

        # Determinar se é negrito/italico no parágrafo inteiro
        if runs_info:
            formatting['is_bold'] = any(r['bold'] for r in runs_info)
            formatting['is_italic'] = any(r['italic'] for r in runs_info)
            formatting['is_underline'] = any(r['underline'] for r in runs_info)

        return formatting

    def read(self) -> Dict[str, Any]:
        """
        Lê o documento preservando a estrutura exata, incluindo:
        - Posição das imagens
        - Listas ordenadas e não ordenadas
        - Hierarquia de títulos
        - Formatação de texto
        """
        # Mapear todas as imagens disponíveis
        image_map: Dict[str, bytes] = {}
        for rel_id, rel in self.doc.part.rels.items():
            if "image" in rel.target_ref:
                try:
                    image_map[rel_id] = rel.target_part.blob
                except Exception:
                    pass

        elements: List[DocumentElement] = []
        seen_images: set = set()

        # Processar cada parágrafo em ordem
        current_list_items: List[ListItem] = []
        current_list_type: Optional[bool] = None  # None = nenhuma lista, True = ordenada, False = não ordenada
        current_list_id = 0

        for paragraph in self.doc.paragraphs:
            p_element = paragraph._p

            # Verificar se há imagens inline neste parágrafo
            paragraph_images = []
            for run in paragraph.runs:
                img = self._extract_image_from_run(run, image_map)
                if img and img.hash not in seen_images:
                    seen_images.add(img.hash)
                    paragraph_images.append(img)

            # Adicionar imagens antes do texto (preservando ordem)
            for img in paragraph_images:
                elements.append(DocumentElement(
                    type='image',
                    content={
                        'data': img.image_data,
                        'filename': img.filename,
                        'hash': img.hash,
                    },
                    formatting={'width': None, 'height': None}
                ))

            # Extrair texto
            text = paragraph.text.strip()

            if not text:
                # Se não tem texto mas tem imagens, já processamos as imagens acima
                continue

            # Verificar se é heading
            heading_level = self._get_heading_level(paragraph)
            if heading_level > 0:
                # Flush lista anterior se existir
                if current_list_items:
                    elements.append(DocumentElement(
                        type='list',
                        content=[{
                            'text': item.text,
                            'level': item.level,
                        } for item in current_list_items],
                        formatting={'ordered': current_list_type}
                    ))
                    current_list_items = []
                    current_list_type = None

                elements.append(DocumentElement(
                    type='heading',
                    content=text,
                    level=heading_level,
                    style=paragraph.style.name if paragraph.style else None,
                    formatting=self._extract_paragraph_formatting(paragraph)
                ))
                continue

            # Verificar se é item de lista
            is_list, is_ordered, list_level = self._is_list_item(paragraph)

            if is_list:
                # Se mudou o tipo de lista, flush a anterior
                if current_list_type is not None and current_list_type != is_ordered:
                    if current_list_items:
                        elements.append(DocumentElement(
                            type='list',
                            content=[{
                                'text': item.text,
                                'level': item.level,
                            } for item in current_list_items],
                            formatting={'ordered': current_list_type}
                        ))
                    current_list_items = []

                current_list_type = is_ordered

                # Limpar prefixos de lista do texto
                clean_text = text
                if is_ordered:
                    clean_text = re.sub(r'^\d+[\.\)]\s*', '', clean_text)
                    clean_text = re.sub(r'^[a-zA-Z][\.\)]\s*', '', clean_text)
                    clean_text = re.sub(r'^[IVXivx]+[\.\)]\s*', '', clean_text)
                else:
                    clean_text = re.sub(r'^[\u2022\u2023\u25E6\u25AA\u25AB\u2013\u2014\-–•◦▪▫–—-]\s*', '', clean_text)
                    clean_text = re.sub(r'^(\[\s*[xX\s]\s*\]|\(\s*[xX\s]\s*\))\s*', '', clean_text)

                current_list_items.append(ListItem(
                    text=clean_text.strip(),
                    level=list_level,
                    is_ordered=is_ordered,
                    numbering_id=current_list_id,
                    list_index=len(current_list_items)
                ))

            else:
                # Não é lista - flush lista anterior se existir
                if current_list_items:
                    elements.append(DocumentElement(
                        type='list',
                        content=[{
                            'text': item.text,
                            'level': item.level,
                        } for item in current_list_items],
                        formatting={'ordered': current_list_type}
                    ))
                    current_list_items = []
                    current_list_type = None

                # Verificar se é subtítulo (heading level 3+)
                style_name = paragraph.style.name if paragraph.style else ""
                is_subheading = False
                sub_level = 0

                if 'Heading' in style_name:
                    try:
                        sub_level = int(style_name.split()[-1])
                        is_subheading = sub_level >= 3
                    except ValueError:
                        pass

                if is_subheading:
                    elements.append(DocumentElement(
                        type='subheading',
                        content=text,
                        level=sub_level,
                        style=style_name,
                        formatting=self._extract_paragraph_formatting(paragraph)
                    ))
                else:
                    # Parágrafo normal
                    elements.append(DocumentElement(
                        type='paragraph',
                        content=text,
                        style=style_name if style_name else None,
                        formatting=self._extract_paragraph_formatting(paragraph)
                    ))

        # Flush lista final se existir
        if current_list_items:
            elements.append(DocumentElement(
                type='list',
                content=[{
                    'text': item.text,
                    'level': item.level,
                } for item in current_list_items],
                formatting={'ordered': current_list_type}
            ))

        # Processar tabelas
        for table in self.doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)

            if table_data:
                elements.append(DocumentElement(
                    type='table',
                    content=table_data,
                ))

        # Preparar lista de imagens para retorno
        images = [
            DocumentImage(
                image_data=el.content['data'],
                filename=el.content['filename'],
                hash=el.content['hash'],
            )
            for el in elements
            if el.type == 'image'
        ]

        return {
            'file_name': self.file_path.name,
            'elements': elements,
            'images': images,
            'total_paragraphs': sum(1 for e in elements if e.type == 'paragraph'),
            'total_headings': sum(1 for e in elements if e.type == 'heading'),
            'total_lists': sum(1 for e in elements if e.type == 'list'),
            'total_tables': sum(1 for e in elements if e.type == 'table'),
            'total_images': len(images),
        }

    def extract_images(self) -> List[DocumentImage]:
        """Extrai todas as imagens do documento (método legado)"""
        images = []

        for rel in self.doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image = rel.target_part
                    image_data = image.blob
                    image_hash = hashlib.md5(image_data).hexdigest()[:8]

                    content_type = image.content_type
                    if "png" in content_type:
                        ext = "png"
                    elif "jpeg" in content_type or "jpg" in content_type:
                        ext = "jpg"
                    elif "gif" in content_type:
                        ext = "gif"
                    else:
                        ext = "png"

                    filename = f"img_{image_hash}.{ext}"

                    images.append(DocumentImage(
                        image_data=image_data,
                        filename=filename,
                        hash=image_hash
                    ))
                except Exception as e:
                    print(f"Erro ao extrair imagem: {e}")

        self.images = images
        return images
