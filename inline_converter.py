#!/usr/bin/env python3
"""
Lector Playbook Converter - Versão com Imagens Inline
Mantém imagens exatamente onde aparecem no texto original
"""

import os
import sys
import re
import io
import hashlib
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass
from typing import List, Dict, Tuple, Optional

sys.path.insert(0, str(Path(__file__).parent / "src"))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    from PIL import Image as PILImage
except ImportError as e:
    print(f"Erro ao importar: {e}")
    os.system("pip install python-docx pillow")
    sys.exit(1)


@dataclass
class BrandIdentity:
    """Identidade visual Lector"""
    primary_dark: str = "#1E3A5F"
    primary_bright: str = "#2563EB"
    accent_coral: str = "#D97757"
    text_dark: str = "#2D3748"
    text_secondary: str = "#4A5568"
    light_bg: str = "#F7FAFC"
    border: str = "#E2E8F0"
    white: str = "#FFFFFF"


class InlineImageExtractor:
    """Extrai imagens e mapeia sua posição no documento"""

    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.images_dir = output_dir / "images"
        self.images_dir.mkdir(parents=True, exist_ok=True)

    def extract_images_with_position(self, doc_path: Path) -> Tuple[Dict[str, Dict], List[Dict]]:
        """
        Extrai imagens e retorna:
        - Map de rel_id -> info da imagem
        - Lista de posições (parágrafo index, rel_id)
        """
        doc = Document(doc_path)
        image_map = {}
        image_positions = []

        # Extrair todas as imagens do documento
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image = rel.target_part
                    image_data = image.blob
                    content_type = image.content_type
                    rel_id = rel.rId

                    # Determinar extensão
                    if "png" in content_type:
                        ext = "png"
                    elif "jpeg" in content_type or "jpg" in content_type:
                        ext = "jpg"
                    elif "gif" in content_type:
                        ext = "gif"
                    else:
                        ext = "png"

                    # Criar hash e nome
                    image_hash = hashlib.md5(image_data).hexdigest()[:8]
                    filename = f"{doc_path.stem}_{image_hash}.{ext}"
                    filepath = self.images_dir / filename

                    # Salvar
                    with open(filepath, "wb") as f:
                        f.write(image_data)

                    # Obter dimensões
                    try:
                        img = PILImage.open(io.BytesIO(image_data))
                        width, height = img.size
                    except:
                        width, height = None, None

                    image_map[rel_id] = {
                        "path": filepath,
                        "filename": filename,
                        "data": image_data,
                        "width": width,
                        "height": height,
                        "extension": ext,
                        "rel_id": rel_id
                    }

                except Exception as e:
                    print(f"    Aviso: Erro ao extrair imagem: {e}")

        # Mapear posições das imagens nos parágrafos
        for para_idx, para in enumerate(doc.paragraphs):
            # Verificar se o parágrafo contém uma imagem
            para_xml = para._p.xml
            for rel_id in image_map.keys():
                if f'r:embed="{rel_id}"' in para_xml or f'r:link="{rel_id}"' in para_xml:
                    image_positions.append({
                        "paragraph_index": para_idx,
                        "rel_id": rel_id,
                        "type": "inline"
                    })

        return image_map, image_positions


@dataclass
class DocumentElement:
    """Representa um elemento do documento"""
    type: str  # 'heading', 'paragraph', 'image', 'table', 'bullet', 'numbered'
    content: any
    level: int = 0
    style: str = ""


class InlinePlaybookConverter:
    """Conversor que mantém imagens inline"""

    def __init__(self, brand: BrandIdentity = None):
        self.brand = brand or BrandIdentity()
        self.image_extractor = None

    def _hex_to_rgb(self, hex_color: str) -> tuple:
        """Converte hex para RGB"""
        hex_color = hex_color.lstrip("#")
        return (
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16)
        )

    def read_docx_with_images(self, file_path: Path) -> Tuple[List[DocumentElement], Dict[str, Dict]]:
        """Lê documento e mapeia elementos com imagens inline"""
        print(f"  Lendo estrutura: {file_path.name}")

        doc = Document(file_path)

        # Extrair imagens com posições
        image_map, image_positions = self.image_extractor.extract_images_with_position(file_path)

        print(f"  Encontradas {len(image_map)} imagem(s) e {len(image_positions)} posições")

        # Criar conjunto de parágrafos que têm imagens
        paragraphs_with_images = {pos["paragraph_index"] for pos in image_positions}

        elements = []
        para_idx = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name if para.style else "Normal"

            # Verificar se este parágrafo tem imagem inline
            if para_idx in paragraphs_with_images:
                # Encontrar imagens neste parágrafo
                para_images = [pos for pos in image_positions if pos["paragraph_index"] == para_idx]

                # Adicionar imagens inline
                for img_pos in para_images:
                    rel_id = img_pos["rel_id"]
                    if rel_id in image_map:
                        elements.append(DocumentElement(
                            type="image",
                            content=image_map[rel_id],
                            level=0,
                            style="inline"
                        ))

            # Adicionar texto do parágrafo (se tiver conteúdo)
            if text:
                if style_name.startswith("Heading"):
                    level = int(style_name.replace("Heading ", "")) if "Heading " in style_name else 2
                    elements.append(DocumentElement(
                        type="heading",
                        content=text,
                        level=level,
                        style=style_name
                    ))
                elif text.isupper() and len(text) < 100:
                    elements.append(DocumentElement(
                        type="heading",
                        content=text,
                        level=2,
                        style="Uppercase"
                    ))
                elif text.startswith(("-", "*", "•")):
                    elements.append(DocumentElement(
                        type="bullet",
                        content=text[1:].strip(),
                        level=0,
                        style="Bullet"
                    ))
                elif re.match(r"^\d+[.\)]\s", text):
                    clean_text = re.sub(r"^\d+[.\)]\s", "", text)
                    elements.append(DocumentElement(
                        type="numbered",
                        content=clean_text,
                        level=0,
                        style="Numbered"
                    ))
                else:
                    elements.append(DocumentElement(
                        type="paragraph",
                        content=text,
                        level=0,
                        style="Normal"
                    ))

            para_idx += 1

        # Adicionar tabelas no final (mantendo ordem aproximada)
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)

            if table_data:
                elements.append(DocumentElement(
                    type="table",
                    content=table_data,
                    level=0,
                    style="Table"
                ))

        return elements, image_map

    def rewrite_text(self, text: str) -> str:
        """Reescreve texto em linguagem mais clara"""
        if not text:
            return text

        replacements = {
            "devera": "deve",
            "deverao": "devem",
            "devera ser": "é",
            "deverao ser": "são",
            "realizar": "fazer",
            "executar": "fazer",
            "inicializar": "abrir",
            "finalizar": "fechar",
            "processo de": "",
            "procedimento de": "",
            "atraves de": "por meio de",
            "no que diz respeito": "sobre",
            "no intuito de": "para",
            "visando": "para",
            "com o objetivo de": "para",
        }

        result = text
        for old, new in replacements.items():
            result = re.sub(r'\b' + old + r'\b', new, result, flags=re.IGNORECASE)

        if result:
            result = result[0].upper() + result[1:]

        return result

    def create_inline_playbook(self, elements: List[DocumentElement], output_path: Path):
        """Cria playbook com imagens inline"""
        print(f"  Criando playbook: {output_path.name}")

        doc = Document()
        self._setup_styles(doc)

        # Header
        self._add_header(doc, output_path.stem.replace("PLAYBOOK_", ""))

        # Processar elementos
        step_counter = 0
        last_was_heading = False

        for element in elements:
            if element.type == "heading":
                level = element.level
                text = self.rewrite_text(element.content)

                if level == 1:
                    heading = doc.add_heading(text, level=1)
                    for run in heading.runs:
                        run.font.size = Pt(20)
                        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_dark))
                    last_was_heading = True
                elif level == 2:
                    heading = doc.add_heading(text, level=2)
                    for run in heading.runs:
                        run.font.size = Pt(16)
                        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_dark))
                    last_was_heading = True
                else:
                    heading = doc.add_heading(text, level=3)
                    for run in heading.runs:
                        run.font.size = Pt(13)
                        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_bright))
                    last_was_heading = True

            elif element.type == "image":
                # Inserir imagem inline exatamente onde está no original
                img_info = element.content
                try:
                    with open(img_info["path"], "rb") as f:
                        image_stream = io.BytesIO(f.read())

                        # Calcular tamanho proporcional
                        if img_info["width"] and img_info["height"]:
                            max_width = Inches(5.5)
                            aspect_ratio = img_info["height"] / img_info["width"]

                            # Limitar altura máxima para imagens altas
                            if aspect_ratio > 1.5:
                                max_width = Inches(3.5)
                        else:
                            max_width = Inches(5)

                        # Adicionar imagem
                        picture = doc.add_picture(image_stream, width=max_width)

                        # Centralizar
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        # Adicionar legenda pequena
                        caption = doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption.add_run(f"[{img_info['filename']}]")
                        caption_run.font.size = Pt(8)
                        caption_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_secondary))
                        caption_run.italic = True

                except Exception as e:
                    print(f"    Aviso: Erro ao inserir imagem: {e}")

                last_was_heading = False

            elif element.type == "bullet":
                text = self.rewrite_text(element.content)
                para = doc.add_paragraph(text, style="List Bullet")
                para.paragraph_format.left_indent = Inches(0.3)
                para.paragraph_format.space_after = Pt(6)
                last_was_heading = False

            elif element.type == "numbered":
                text = self.rewrite_text(element.content)
                para = doc.add_paragraph(text, style="List Number")
                para.paragraph_format.left_indent = Inches(0.3)
                para.paragraph_format.space_after = Pt(6)
                last_was_heading = False

            elif element.type == "paragraph":
                text = self.rewrite_text(element.content)

                if self._is_procedure_step(text):
                    step_counter += 1
                    self._add_step(doc, step_counter, text)
                else:
                    para = doc.add_paragraph(text)
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_after = Pt(8)

                last_was_heading = False

            elif element.type == "table":
                self._add_table(doc, element.content)
                last_was_heading = False

        # Footer
        self._add_footer(doc)

        # Salvar
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        print(f"  Salvo: {output_path}")

    def _setup_styles(self, doc):
        """Configura estilos"""
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_dark))

    def _add_header(self, doc, title: str):
        """Adiciona header"""
        doc.add_paragraph()

        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run("PLAYBOOK")
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_bright))

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title.upper())
        title_run.font.size = Pt(26)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_dark))

        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line.add_run("━" * 35)
        line_run.font.size = Pt(8)
        line_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_bright))

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run("Guia Prático | Lector Tecnologia")
        sub_run.font.size = Pt(10)
        sub_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_secondary))
        sub_run.italic = True

        doc.add_paragraph()

        about = doc.add_paragraph()
        about_run = about.add_run("SOBRE ESTE PLAYBOOK")
        about_run.font.size = Pt(12)
        about_run.font.bold = True
        about_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_dark))

        about_text = doc.add_paragraph()
        about_text_run = about_text.add_run(
            "Este guia foi desenvolvido para ser prático, direto e fácil de seguir. "
            "As imagens aparecem exatamente onde são referenciadas no texto original."
        )
        about_text_run.font.size = Pt(10)
        about_text.paragraph_format.left_indent = Inches(0.2)

        doc.add_paragraph()

    def _add_step(self, doc, number: int, text: str):
        """Adiciona passo"""
        step = doc.add_paragraph()
        step.paragraph_format.left_indent = Inches(0.2)

        num_run = step.add_run(f"PASSO {number:02d} ")
        num_run.font.size = Pt(11)
        num_run.font.bold = True
        num_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_bright))

        text_run = step.add_run(text)
        text_run.font.size = Pt(11)
        text_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_dark))

        doc.add_paragraph()

    def _add_table(self, doc, data: List[List[str]]):
        """Adiciona tabela"""
        if not data or len(data) < 1:
            return

        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = "Table Grid"

        for i, row in enumerate(data):
            for j, cell_text in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = self.rewrite_text(cell_text)

                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)

                    shading_elm = parse_xml(
                        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                        f'w:fill="{self.brand.primary_dark.lstrip("#")}"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading_elm)

        doc.add_paragraph()

    def _add_footer(self, doc):
        """Adiciona footer"""
        doc.add_page_break()

        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line.add_run("━" * 35)
        line_run.font.size = Pt(8)
        line_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.border))

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("LECTOR TECNOLOGIA")
        footer_run.font.size = Pt(12)
        footer_run.font.bold = True
        footer_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.primary_dark))

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run("www.lector.com.br")
        sub_run.font.size = Pt(9)
        sub_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.brand.text_secondary))

    def _is_procedure_step(self, text: str) -> bool:
        """Detecta passo de procedimento"""
        indicators = [
            r"\b(clique|selecione|digite|preencha|envie|acesse|abra|faça)\b",
            r"\b(proximo passo|em seguida|depois)\b",
            r"\b(insira|adicione|remova|exclua|atualize)\b",
        ]

        for pattern in indicators:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False


def main():
    print("=" * 65)
    print("  LECTOR PLAYBOOK CONVERTER")
    print("  Imagens Inline (exatamente onde aparecem no original)")
    print("=" * 65)

    input_dir = Path(r"C:\Users\Lector\Desktop\Documentações")
    output_dir = Path(r"C:\Users\Lector\lector-playbook-converter\output")

    if not input_dir.exists():
        print(f"\nErro: Pasta não encontrada: {input_dir}")
        return 1

    docx_files = list(input_dir.glob("*.docx"))
    if not docx_files:
        print(f"\nNenhum arquivo .docx encontrado")
        return 1

    print(f"\nEncontrados {len(docx_files)} documento(s):")
    for f in docx_files:
        print(f"  - {f.name}")

    brand = BrandIdentity()
    converter = InlinePlaybookConverter(brand)
    converter.image_extractor = InlineImageExtractor(output_dir)

    print("\n" + "-" * 65)
    results = []

    for docx_file in docx_files:
        try:
            print(f"\n>> Processando: {docx_file.name}")
            print("-" * 50)

            # Ler com imagens inline
            elements, image_map = converter.read_docx_with_images(docx_file)

            print(f"  Total de elementos: {len(elements)}")

            # Criar playbook
            output_file = output_dir / f"PLAYBOOK_{docx_file.stem}_INLINE.docx"
            converter.create_inline_playbook(elements, output_file)

            img_count = sum(1 for e in elements if e.type == "image")
            results.append((docx_file.name, output_file.name, img_count, "Sucesso"))
            print(f"  OK: {output_file.name}")

        except Exception as e:
            print(f"  ERRO: {e}")
            import traceback
            traceback.print_exc()
            results.append((docx_file.name, None, 0, f"Erro: {e}"))

    print("\n" + "=" * 65)
    print("  RESUMO")
    print("=" * 65)

    for input_name, output_name, img_count, status in results:
        if status == "Sucesso":
            print(f"  OK: {input_name}")
            print(f"      -> {output_name}")
            print(f"      -> {img_count} imagem(ns) mantida(s) inline")
        else:
            print(f"  ERRO: {input_name} - {status}")

    print(f"\n  Pasta: {output_dir}")
    print("=" * 65 + "\n")

    return 0


if __name__ == "__main__":
    sys.exit(main())
